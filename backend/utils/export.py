"""Export utilities for chat conversations."""

import json
from datetime import datetime
from io import BytesIO
from pathlib import Path
from typing import List

from docx import Document as DocxDocument
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer

from backend.database.models import Conversation, Message


def export_to_markdown(conversation: Conversation) -> str:
    lines = [
        f"# {conversation.title}",
        f"*Exported on {datetime.now().strftime('%Y-%m-%d %H:%M')}*",
        "",
    ]

    for msg in conversation.messages:
        role = "You" if msg.role == "user" else "ResearchGPT"
        lines.append(f"## {role}")
        lines.append(msg.content)
        lines.append("")

        if msg.citations:
            try:
                citations = json.loads(msg.citations)
                if citations:
                    lines.append("### Sources")
                    for c in citations:
                        lines.append(
                            f"- **{c['document_name']}** (Page {c['page_number']}) — "
                            f"Confidence: {c['confidence_score']}%"
                        )
                    lines.append("")
            except (json.JSONDecodeError, TypeError):
                pass

    return "\n".join(lines)


def export_to_docx(conversation: Conversation) -> BytesIO:
    doc = DocxDocument()
    doc.add_heading(conversation.title, 0)
    doc.add_paragraph(f"Exported on {datetime.now().strftime('%Y-%m-%d %H:%M')}")

    for msg in conversation.messages:
        role = "You" if msg.role == "user" else "ResearchGPT"
        doc.add_heading(role, level=2)
        doc.add_paragraph(msg.content)

        if msg.citations:
            try:
                citations = json.loads(msg.citations)
                if citations:
                    doc.add_heading("Sources", level=3)
                    for c in citations:
                        doc.add_paragraph(
                            f"{c['document_name']} (Page {c['page_number']}) — "
                            f"Confidence: {c['confidence_score']}%"
                        )
            except (json.JSONDecodeError, TypeError):
                pass

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


def export_to_pdf(conversation: Conversation) -> BytesIO:
    buffer = BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=letter)
    styles = getSampleStyleSheet()
    story: List = []

    story.append(Paragraph(conversation.title, styles["Title"]))
    story.append(Spacer(1, 12))

    for msg in conversation.messages:
        role = "You" if msg.role == "user" else "ResearchGPT"
        story.append(Paragraph(f"<b>{role}</b>", styles["Heading2"]))
        content = msg.content.replace("\n", "<br/>")
        story.append(Paragraph(content, styles["Normal"]))
        story.append(Spacer(1, 12))

    doc.build(story)
    buffer.seek(0)
    return buffer
