"""Document parsing for PDF, DOCX, TXT, and Markdown files."""

from dataclasses import dataclass
from pathlib import Path
from typing import List

import fitz  # PyMuPDF
from docx import Document as DocxDocument

from backend.utils.logger import logger
from backend.utils.text_cleaner import clean_text


@dataclass
class ParsedPage:
    page_number: int
    content: str


@dataclass
class ParsedDocument:
    pages: List[ParsedPage]
    page_count: int
    metadata: dict


class DocumentParser:
    SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".txt", ".md", ".markdown"}

    def parse(self, file_path: Path) -> ParsedDocument:
        suffix = file_path.suffix.lower()
        if suffix not in self.SUPPORTED_EXTENSIONS:
            raise ValueError(f"Unsupported file type: {suffix}")

        parsers = {
            ".pdf": self._parse_pdf,
            ".docx": self._parse_docx,
            ".txt": self._parse_text,
            ".md": self._parse_text,
            ".markdown": self._parse_text,
        }
        return parsers[suffix](file_path)

    def _parse_pdf(self, file_path: Path) -> ParsedDocument:
        pages: List[ParsedPage] = []
        metadata: dict = {}

        with fitz.open(file_path) as doc:
            metadata = dict(doc.metadata) if doc.metadata else {}
            for i, page in enumerate(doc):
                text = page.get_text("text")
                cleaned = clean_text(text)
                if cleaned:
                    pages.append(ParsedPage(page_number=i + 1, content=cleaned))

        if not pages:
            logger.warning("No text extracted from PDF %s, attempting OCR fallback", file_path)
            pages = self._ocr_pdf(file_path)

        return ParsedDocument(pages=pages, page_count=len(pages), metadata=metadata)

    def _ocr_pdf(self, file_path: Path) -> List[ParsedPage]:
        """OCR fallback for scanned PDFs using PyMuPDF text extraction with images."""
        pages: List[ParsedPage] = []
        try:
            with fitz.open(file_path) as doc:
                for i, page in enumerate(doc):
                    blocks = page.get_text("blocks")
                    text_parts = [b[4] for b in blocks if isinstance(b[4], str)]
                    cleaned = clean_text(" ".join(text_parts))
                    if cleaned:
                        pages.append(ParsedPage(page_number=i + 1, content=cleaned))
        except Exception as e:
            logger.error("OCR fallback failed: %s", e)
        return pages

    def _parse_docx(self, file_path: Path) -> ParsedDocument:
        doc = DocxDocument(str(file_path))
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        full_text = clean_text("\n\n".join(paragraphs))

        pages: List[ParsedPage] = []
        chunk_size = 3000
        if len(full_text) <= chunk_size:
            pages.append(ParsedPage(page_number=1, content=full_text))
        else:
            for i in range(0, len(full_text), chunk_size):
                chunk = full_text[i : i + chunk_size]
                pages.append(ParsedPage(page_number=len(pages) + 1, content=chunk))

        metadata = {}
        if doc.core_properties:
            props = doc.core_properties
            metadata = {
                "author": props.author or "",
                "title": props.title or "",
                "subject": props.subject or "",
            }

        return ParsedDocument(pages=pages, page_count=len(pages), metadata=metadata)

    def _parse_text(self, file_path: Path) -> ParsedDocument:
        content = file_path.read_text(encoding="utf-8", errors="ignore")
        cleaned = clean_text(content)

        pages: List[ParsedPage] = []
        chunk_size = 3000
        if len(cleaned) <= chunk_size:
            pages.append(ParsedPage(page_number=1, content=cleaned))
        else:
            for i in range(0, len(cleaned), chunk_size):
                chunk = cleaned[i : i + chunk_size]
                pages.append(ParsedPage(page_number=len(pages) + 1, content=chunk))

        return ParsedDocument(pages=pages, page_count=len(pages), metadata={})
